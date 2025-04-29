from flask import Blueprint, request, jsonify, send_file
from sqlalchemy import or_, and_
from flask_jwt_extended import jwt_required, get_jwt_identity
from app.utils import admin_required
from app import db
from app.models import Template, StudentCourse, Student, CourseStatus, OperationLog
from datetime import datetime
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

bp = Blueprint('templates', __name__, url_prefix='/api/v1/templates')

@bp.route('/report_card/<int:sc_id>', methods=['POST'])
@jwt_required()
def generate_report_card(sc_id):
    """生成并发布成绩单"""
    try:
        sc = StudentCourse.query.get_or_404(sc_id)
        current_user_id = get_jwt_identity()

        # 检查是否已有期中期末成绩
        if sc.midterm_grade is None or sc.final_grade is None:
            return jsonify({'code': 400, 'message': '必须同时有期中成绩和期末成绩才能生成成绩单'}), 400

        # 获取成绩单模板
        template = Template.query.filter_by(template_type='report_card').first()
        if not template:
            return jsonify({'code': 404, 'message': '未找到成绩单模板'}), 404

        # 记录开始生成成绩单日志
        log = OperationLog(
            user_id=current_user_id,
            operation_type='start_generate_report_card',
            operation_detail=f'开始为学生 {sc.student.first_name} {sc.student.last_name} 生成课程 {sc.course_code} 的成绩单',
            target_table='student_courses',
            target_id=sc.id,
            created_at=datetime.now()
        )
        db.session.add(log)

        # 生成成绩单文档
        doc = Document()
        
        # 添加标题
        title = doc.add_paragraph('成绩单')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(16)
        title.runs[0].font.bold = True
        
        # 添加学生信息
        doc.add_paragraph(f'学生姓名: {sc.student.first_name} {sc.student.last_name}')
        doc.add_paragraph(f'课程代码: {sc.course_code}')
        doc.add_paragraph(f'课程名称: {sc.course.name}')
        doc.add_paragraph(f'期中成绩: {sc.midterm_grade}')
        doc.add_paragraph(f'期末成绩: {sc.final_grade}')
        
        # 计算总评成绩
        total_grade = (sc.midterm_grade + sc.final_grade) / 2
        doc.add_paragraph(f'总评成绩: {total_grade}')
        
        # 添加日期和签名
        doc.add_paragraph(f'日期: {datetime.now().strftime("%Y-%m-%d")}')
        doc.add_paragraph('教师签名: _________________')
        
        # 保存文档
        file_path = f'report_cards/{sc.student_id}_{sc.course_code}_{datetime.now().strftime("%Y%m%d")}.docx'
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        doc.save(file_path)

        # 更新课程状态为已修
        sc.status = CourseStatus.GRADUATED
        sc.completion_date = datetime.now().date()
        sc.report_card_date = datetime.now().date()

        # 记录发布成绩单日志
        log = OperationLog(
            user_id=current_user_id,
            operation_type='publish_report_card',
            operation_detail=f'发布学生 {sc.student.first_name} {sc.student.last_name} 的课程 {sc.course_code} 成绩单，期中成绩 {sc.midterm_grade}，期末成绩 {sc.final_grade}',
            target_table='student_courses',
            target_id=sc.id,
            created_at=datetime.now()
        )
        db.session.add(log)

        db.session.commit()

        return jsonify({
            'code': 200,
            'data': {
                'file_path': file_path,
                'student_course': sc.to_dict()
            },
            'message': '成绩单生成并发布成功'
        }), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({'code': 500, 'message': f'生成成绩单失败: {str(e)}'}), 500

@bp.route('/report_card/<int:sc_id>/download', methods=['GET'])
@jwt_required()
def download_report_card(sc_id):
    """下载成绩单"""
    try:
        sc = StudentCourse.query.get_or_404(sc_id)
        current_user_id = get_jwt_identity()
        
        file_path = f'report_cards/{sc.student_id}_{sc.course_code}_{sc.report_card_date.strftime("%Y%m%d")}.docx'
        
        if not os.path.exists(file_path):
            return jsonify({'code': 404, 'message': '成绩单文件不存在'}), 404
            
        # 记录下载成绩单日志
        log = OperationLog(
            user_id=current_user_id,
            operation_type='download_report_card',
            operation_detail=f'下载学生 {sc.student.first_name} {sc.student.last_name} 的课程 {sc.course_code} 成绩单',
            target_table='student_courses',
            target_id=sc.id,
            created_at=datetime.now()
        )
        db.session.add(log)
        db.session.commit()
            
        return send_file(file_path, as_attachment=True, download_name=f'report_card_{sc.course_code}.docx')

    except Exception as e:
        return jsonify({'code': 500, 'message': f'下载成绩单失败: {str(e)}'}), 500 